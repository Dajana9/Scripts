from os import listdir
from os.path import isfile, join, isdir
import sys
import webbrowser as wb
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from docx import Document
import PyPDF2


class App(QWidget):
 
    def __init__(self):
        super().__init__()
        self.title = 'Hey =D'
        self.left = 10
        self.top = 10
        self.width = 640
        self.height = 480
        self.initUI()
 
    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)
        
        self.textbox_label()
        self.pushbutton()
        self.listWidget()
        self.layouts()

        self.show()

    def layouts(self):
        grid = QGridLayout()
        self.setLayout(grid)
        grid.setSpacing(15)

        grid.addWidget(self.label_intro, 1, 0, 1, 2)
        grid.addWidget(self.label_dir, 2, 0)
        grid.addWidget(self.textbox_dir, 2, 1)
        grid.addWidget(self.button, 2, 3)

        grid.addWidget(self.label_words, 3, 0)
        grid.addWidget(self.textbox_words, 3, 1)

        grid.addWidget(self.label_list, 4, 0, 1, 2)
        grid.addWidget(self.listWidget, 5, 0, 5, 2)

    def listWidget(self):   
        self.label_list = QLabel('Double click on document if you want to open it!', self)
        self.listWidget = QListWidget(self)
        self.listWidget.itemDoubleClicked.connect(self.open_file)

    def open_file(self, item):
        wb.open_new(item.text())

    def pushbutton(self):
        self.button = QPushButton('Search', self)
        self.button.clicked.connect(self.on_click)

    def textbox_label(self):
        self.label_intro = QLabel("""Hey, if you are looking for a document that is containing some terms,
    but forgot the name of the document, maybe this can help!\n
            Enter the path to the directory from which you want to start search.
            Enter the words you are looking for.
            The program will enter all subdirectories in the search...""", self)

        self.label_intro.setFont(QFont("Times",weight=QFont.Bold))

        self.label_dir = QLabel('Directory:', self)
        self.textbox_dir = QLineEdit(self)

        self.label_words = QLabel('Words:', self)
        self.textbox_words = QLineEdit(self)

    def on_click(self):
        directory = self.textbox_dir.text()
        self.listWidget.clear()    
        self.go_trough_all_dir(directory)    

    def directory_not_found(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Information)
        msg.setText("The directory you are looking for doesn't exist")
        msg.setInformativeText("You must write the complete path to you directory, for example try /home/{USER}/Documents..")
        msg.setWindowTitle("Directory not found")
        msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)

        retval = msg.exec_()

    def go_trough_all_dir(self, directory):
        try:
            for file_iterator in listdir(directory):
                file_path = join(directory, file_iterator)
                self.search_words(file_path)
                if isdir(file_path):
                    self.go_trough_all_dir(file_path)
        except FileNotFoundError:
            self.directory_not_found()

    def search_words(self, file_path):
        if isfile(file_path):
            are_all_words_in_file = True
            word_list = self.textbox_words.text().split(',')
            for word in word_list:
                if file_path[-3:] == 'pdf':
                        pdfFileObj = open(file_path, 'rb')
                        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) #za citanje word i pdf prilagodi..try catch opcije
                        for page_num in range(pdfReader.numPages): 
                            if word not in pdfReader.getPage(page_num).extractText():
                                are_all_words_in_file = False
                                break
                elif file_path[-4:] == 'docx':
                    document = Document(file_path)
                    for paragraph in document.paragraphs:
                        if word not in paragraph.text:
                            #the word is not in paragraphs, check the tables now
                            for table in document.tables:
                                for cell in table.cells:
                                    for paragraph in cell.paragraphs:
                                        if word not in paragraph.text:
                                            are_all_words_in_file = False
                                            break
                else:   
                    try:
                        if word not in open(file_path).read():  
                            are_all_words_in_file = False
                            break
                    except UnicodeDecodeError:
                        are_all_words_in_file = False
                        break
            if are_all_words_in_file is True:             
                item = QListWidgetItem(file_path)
                self.listWidget.addItem(item)

def main():
    app = QApplication(sys.argv)
    ex = App()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()